"""Document parser and text chunking engine for PDF, DOCX, TXT, MD, CSV, SQL DB (.db, .sqlite, .sql), and NoSQL (.json, .jsonl) files."""

from __future__ import annotations
import hashlib
import json
import os
import re
import sqlite3
import tempfile
from io import BytesIO, StringIO
from pathlib import Path


def file_hash(data: bytes) -> str:
    """Return SHA-256 hash of file content for deduplication."""
    return hashlib.sha256(data).hexdigest()


def extract_text_from_pdf(stream: BytesIO) -> list[dict]:
    """Extract pages from PDF stream preserving tabular layout (Grades, Marks, Tables)."""
    pages = []
    # Try pypdf with layout mode first for exact tabular alignment
    try:
        import pypdf
        reader = pypdf.PdfReader(stream)
        for idx, page in enumerate(reader.pages, start=1):
            try:
                text = page.extract_text(extraction_mode="layout") or ""
            except Exception:
                text = page.extract_text() or ""
            text = text.strip()
            if text:
                pages.append({"page": idx, "text": text})
        if pages:
            return pages
    except Exception:
        pass

    # Fallback to PyPDF2
    try:
        import PyPDF2
        reader = PyPDF2.PdfReader(stream)
        for idx, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            text = text.strip()
            if text:
                pages.append({"page": idx, "text": text})
    except Exception as e:
        print(f"Error reading PDF: {e}")
    return pages


def extract_text_from_docx(stream: BytesIO) -> list[dict]:
    """Extract paragraphs and tables from DOCX stream."""
    try:
        import docx
        doc = docx.Document(stream)
        text_blocks = []
        for p in doc.paragraphs:
            if p.text.strip():
                text_blocks.append(p.text.strip())
        
        # Extract tables as well
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                if row_text:
                    text_blocks.append(row_text)

        full_text = "\n\n".join(text_blocks)
        if full_text:
            return [{"page": 1, "text": full_text}]
    except Exception as e:
        print(f"Error reading DOCX: {e}")
    return []


def extract_text_from_sqlite(stream: BytesIO, filename: str) -> list[dict]:
    """Extract schema and data rows from a SQLite relational database file (.db, .sqlite)."""
    pages = []
    tmp_path = None
    try:
        with tempfile.NamedTemporaryFile(suffix=".db", delete=False) as tmp:
            tmp.write(stream.getvalue())
            tmp_path = tmp.name

        conn = sqlite3.connect(tmp_path)
        cursor = conn.cursor()

        # Get all table names
        cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name NOT LIKE 'sqlite_%';")
        tables = [row[0] for row in cursor.fetchall()]

        for idx, table_name in enumerate(tables, start=1):
            cursor.execute(f"PRAGMA table_info('{table_name}');")
            columns = [col[1] for col in cursor.fetchall()]
            
            cursor.execute(f"SELECT * FROM '{table_name}';")
            rows = cursor.fetchall()

            table_lines = [
                f"[SQL Database: {filename} | Table: {table_name}]",
                f"Columns: {', '.join(columns)}",
                f"Total Rows: {len(rows)}",
                "--- DATA ROWS ---"
            ]

            for row_idx, row in enumerate(rows, start=1):
                row_pairs = [f"{col}: {val}" for col, val in zip(columns, row) if val is not None]
                table_lines.append(f"Row {row_idx} -> " + " | ".join(row_pairs))

            table_text = "\n".join(table_lines)
            pages.append({"page": idx, "text": table_text})

        conn.close()
    except Exception as e:
        print(f"Error reading SQLite database '{filename}': {e}")
    finally:
        if tmp_path and os.path.exists(tmp_path):
            try:
                os.remove(tmp_path)
            except Exception:
                pass

    return pages


def extract_text_from_sql_dump(stream: BytesIO, filename: str) -> list[dict]:
    """Extract tables, schema DDL, and INSERT records from a raw .sql dump script."""
    try:
        raw_sql = stream.read().decode("utf-8", errors="ignore").strip()
        lines = [line.strip() for line in raw_sql.splitlines() if line.strip() and not line.startswith("--")]
        clean_text = "\n".join(lines)
        if clean_text:
            header = f"[SQL Database Script: {filename}]\n"
            return [{"page": 1, "text": header + clean_text}]
    except Exception as e:
        print(f"Error reading SQL dump '{filename}': {e}")
    return []


def extract_text_from_nosql_json(stream: BytesIO, filename: str) -> list[dict]:
    """Extract document collections and key-value records from NoSQL JSON / JSONL files."""
    pages = []
    try:
        raw_content = stream.read().decode("utf-8", errors="ignore").strip()
        data = None

        # Try standard JSON first
        try:
            data = json.loads(raw_content)
        except Exception:
            # Fallback to JSONL (line-delimited JSON)
            data = []
            for line in raw_content.splitlines():
                line = line.strip()
                if line:
                    try:
                        data.append(json.loads(line))
                    except Exception:
                        pass

        if isinstance(data, list):
            doc_lines = [f"[NoSQL Collection: {filename} | Total Documents: {len(data)}]"]
            for idx, doc in enumerate(data, start=1):
                if isinstance(doc, dict):
                    fields = [f"{k}: {v}" for k, v in doc.items()]
                    doc_lines.append(f"Document #{idx} -> " + " | ".join(fields))
                else:
                    doc_lines.append(f"Document #{idx} -> {doc}")
            pages.append({"page": 1, "text": "\n".join(doc_lines)})

        elif isinstance(data, dict):
            doc_lines = [f"[NoSQL Document Store: {filename}]"]
            for key, val in data.items():
                if isinstance(val, list):
                    doc_lines.append(f"Collection '{key}' ({len(val)} items):")
                    for item_idx, item in enumerate(val, start=1):
                        if isinstance(item, dict):
                            fields = [f"{k}: {v}" for k, v in item.items()]
                            doc_lines.append(f"  Item #{item_idx} -> " + " | ".join(fields))
                        else:
                            doc_lines.append(f"  Item #{item_idx} -> {item}")
                else:
                    doc_lines.append(f"{key}: {val}")
            pages.append({"page": 1, "text": "\n".join(doc_lines)})

        else:
            pages.append({"page": 1, "text": f"[NoSQL Data: {filename}]\n{raw_content}"})

    except Exception as e:
        print(f"Error reading NoSQL file '{filename}': {e}")
    return pages


def extract_text_from_excel(stream: BytesIO, filename: str) -> list[dict]:
    """Universal, domain-agnostic CSV-preprocessed parser for ANY Excel spreadsheet (.xlsx, .xls, .xlsm)."""
    pages = []
    # Primary: pandas/openpyxl CSV preprocessing with Categorical Overview Header
    try:
        import pandas as pd
        excel_file = pd.ExcelFile(stream)
        for idx, sheet_name in enumerate(excel_file.sheet_names, start=1):
            df_sheet = pd.read_excel(excel_file, sheet_name=sheet_name).dropna(how='all')
            if df_sheet.empty:
                continue

            headers = [str(c).strip() for c in df_sheet.columns]
            col_value_counts = {h: {} for h in headers}

            # Gather categorical distribution statistics across all rows
            for _, row in df_sheet.iterrows():
                for h in headers:
                    val = str(row[h]).strip() if pd.notna(row[h]) else ""
                    if val != "" and len(val) <= 50:
                        col_value_counts[h][val] = col_value_counts[h].get(val, 0) + 1

            # Build Categorical Column Distribution Summary
            summary_lines = []
            cat_summaries = []
            for h in headers:
                counts = col_value_counts[h]
                if 2 <= len(counts) <= 100 and sum(counts.values()) >= 4:
                    top_vals = [f"'{val}' ({cnt})" for val, cnt in sorted(counts.items(), key=lambda x: x[1], reverse=True)]
                    cat_summaries.append(f"- Column [{h}]: {', '.join(top_vals)}")

            if cat_summaries:
                summary_lines.append("=== Sheet Overview & Categorical Distributions ===")
                summary_lines.extend(cat_summaries)
                summary_lines.append("\n=== CSV Records ===")

            # Convert sheet rows to clean, space-efficient CSV
            csv_buffer = StringIO()
            df_sheet.to_csv(csv_buffer, index=False)
            csv_records_str = csv_buffer.getvalue()

            sheet_header = f"[Excel Dataset: {filename} | Sheet: {sheet_name} | Total Columns: {len(headers)} | Total Rows: {len(df_sheet)}]\n"
            full_sheet_text = sheet_header + "\n".join(summary_lines) + "\n" + csv_records_str
            pages.append({"page": idx, "text": full_sheet_text})

        if pages:
            return pages
    except Exception as e:
        print(f"CSV preprocessed Excel extraction for '{filename}' failed: {e}")

    # Fallback: openpyxl raw string extraction
    try:
        import openpyxl
        wb = openpyxl.load_workbook(stream, data_only=True)
        for idx, sheet_name in enumerate(wb.sheetnames, start=1):
            sheet = wb[sheet_name]
            rows_text = []
            for row in sheet.iter_rows(values_only=True):
                cell_vals = [str(val).strip() for val in row if val is not None and str(val).strip() != ""]
                if cell_vals:
                    rows_text.append(",".join(cell_vals))
            if rows_text:
                pages.append({"page": idx, "text": f"[Excel Workbook: {filename} | Sheet: {sheet_name}]\n" + "\n".join(rows_text)})
        if pages:
            return pages
    except Exception as e:
        print(f"openpyxl fallback for '{filename}' failed: {e}")

    return []


def extract_text_from_txt(stream: BytesIO) -> list[dict]:
    """Extract text from plain text or markdown stream with binary guard."""
    try:
        raw_bytes = stream.read()
        # Binary guard: if content contains null bytes or zip headers, do not parse as text
        if b"\x00" in raw_bytes[:1024] or raw_bytes.startswith(b"PK\x03\x04"):
            print("Binary file detected in text parser. Skipping binary content.")
            return []
        
        text = raw_bytes.decode("utf-8", errors="ignore").strip()
        if text:
            return [{"page": 1, "text": text}]
    except Exception as e:
        print(f"Error reading text file: {e}")
    return []


def extract_pages(stream: BytesIO, filename: str) -> list[dict]:
    """Extract pages/text blocks based on file extension, including Excel, SQL, and NoSQL databases."""
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(stream)
    elif ext in [".docx", ".doc"]:
        return extract_text_from_docx(stream)
    elif ext in [".xlsx", ".xls", ".xlsm"]:
        return extract_text_from_excel(stream, filename)
    elif ext in [".db", ".sqlite", ".sqlite3"]:
        return extract_text_from_sqlite(stream, filename)
    elif ext == ".sql":
        return extract_text_from_sql_dump(stream, filename)
    elif ext in [".json", ".jsonl"]:
        return extract_text_from_nosql_json(stream, filename)
    elif ext in [".txt", ".md", ".csv"]:
        return extract_text_from_txt(stream)
    else:
        return extract_text_from_txt(stream)


def chunk_pages(pages: list[dict], source_name: str, chunk_size: int = 500, chunk_overlap: int = 50) -> list[dict]:
    """Split extracted pages into overlapping text chunks with rich metadata and header prefix preservation."""
    chunks = []
    chunk_counter = 0

    for page_info in pages:
        page_num = page_info["page"]
        text = page_info["text"]

        # Check if text has a structured Overview/Header prefix (for Excel, SQL, NoSQL)
        header_prefix = ""
        body_text = text

        if "=== Detailed Records ===" in text:
            parts = text.split("=== Detailed Records ===")
            header_prefix = parts[0].strip() + "\n=== Detailed Records ===\n"
            body_text = parts[1].strip()
        elif "\nRow #" in text:
            lines = text.split("\nRow #")[0]
            header_prefix = lines.strip() + "\n"
            body_text = text[len(lines):].strip()
        elif "\nDocument #" in text:
            lines = text.split("\nDocument #")[0]
            header_prefix = lines.strip() + "\n"
            body_text = text[len(lines):].strip()

        words = body_text.split()
        prefix_words = header_prefix.split() if header_prefix else []
        effective_chunk_size = max(100, chunk_size - len(prefix_words))
        step = max(50, effective_chunk_size - chunk_overlap)

        if len(words) <= effective_chunk_size:
            chunk_text = (header_prefix + body_text).strip()
            chunk_id = f"{source_name}_p{page_num}_c{chunk_counter}"
            chunks.append({
                "id": chunk_id,
                "text": chunk_text,
                "metadata": {
                    "source": source_name,
                    "page": page_num,
                    "chunk_index": chunk_counter,
                    "word_count": len(words) + len(prefix_words)
                }
            })
            chunk_counter += 1
        else:
            for i in range(0, len(words), step):
                chunk_words = words[i:i + effective_chunk_size]
                if not chunk_words:
                    continue
                sub_body = " ".join(chunk_words)
                chunk_text = (header_prefix + sub_body).strip()
                chunk_id = f"{source_name}_p{page_num}_c{chunk_counter}"
                chunks.append({
                    "id": chunk_id,
                    "text": chunk_text,
                    "metadata": {
                        "source": source_name,
                        "page": page_num,
                        "chunk_index": chunk_counter,
                        "word_count": len(chunk_words) + len(prefix_words)
                    }
                })
                chunk_counter += 1

    return chunks
